from tkinter import *
from tkinter import ttk
import sqlite3
import os
import time
import calendar
import datetime

from docx import Document


class ReportObject(Frame):

    db_directory = 'ReportsDatabase'

    if not os.path.exists(db_directory):
        os.makedirs(db_directory)

    db_name = db_directory + '/' + 'reportdb.db'

    def __init__(self, root):

        super().__init__()
        self.root = root
        self.root.title('Reports')
        self.initial_migration()
        # self.viewing_records()

        labelsFrameSave = LabelFrame(root, text='Запись отчета в базу данных')
        labelsFrameSave.grid(column=0, row=1)

        labelsFrameDownload = LabelFrame(root, text='Выгрузка отчета в формате .docx')
        labelsFrameDownload.grid(column=0, row=2)

        Label(labelsFrameSave, text='Дата:').grid(row=1, column=1, sticky=W)
        self.event_date = Entry(labelsFrameSave)
        self.event_date.grid(row=1, column=2)

        self.event_date.insert(0, str(datetime.date.today()))

        self.date_entry = Button(labelsFrameSave, text='Выбрать', command=DatePickEventDate)
        self.date_entry.grid(row=1, column=3)

        Label(labelsFrameSave, text='Уровень:').grid(row=2, column=1, sticky=W)
        self.level_box_value = StringVar()
        self.level = ttk.Combobox(labelsFrameSave, textvariable=self.level_box_value)
        self.level['values'] = (
            'HIGH',
            'MID',
            'LOW'
        )
        self.level.grid(row=2, column=2)

        Label(labelsFrameSave, text='Событие:').grid(row=3, column=1, sticky=W)
        self.event_box_value = StringVar()
        self.event = ttk.Combobox(labelsFrameSave, textvariable=self.event_box_value)
        self.event['values'] = (
            'Suspicious user-agent detected',
            'Suspicious Behaviour',
            'Worm infection',
            'AV Attacks',
            'AV Malware',
            'Malware Infection',
            'Keitaro TDS'
        )
        self.event.grid(row=3, column=2)

        Label(labelsFrameSave, text='Срабатываний:').grid(row=4, column=1, sticky=W)
        self.quantity = Entry(labelsFrameSave)
        self.quantity.grid(row=4, column=2)

        Label(labelsFrameSave, text='Статус:').grid(row=5, column=1, sticky=W)
        self.status_box_value = StringVar()
        self.status = ttk.Combobox(labelsFrameSave, textvariable=self.status_box_value)
        self.status['values'] = (
            'Ложный',
            'Реальный'
        )
        self.status.grid(row=5, column=2)

        Label(labelsFrameSave, text='Коментарий:').grid(row=6, column=1, sticky=NW)
        self.comment = Text(labelsFrameSave, height=5, width=20)
        self.comment.grid(row=6, column=2)

        ttk.Button(labelsFrameSave, text='Сохранить в БД', command=self.adding).grid(row=7, column=2)
        self.message = Label(labelsFrameSave, text='', fg='green')
        self.message.grid(row=8, column=1)

        # Блок выгрузки .docx
        Label(labelsFrameDownload, text='Дата c:').grid(row=9, column=1)
        self.date_from = Entry(labelsFrameDownload)
        self.date_from.grid(row=9, column=2)

        self.date_from_btn = Button(labelsFrameDownload, text='Выбрать', command=DatePickReportDateFrom)
        self.date_from_btn.grid(row=9, column=3)

        Label(labelsFrameDownload, text='Дата по:').grid(row=10, column=1)

        self.date_to = Entry(labelsFrameDownload)
        self.date_to.grid(row=10, column=2)

        self.date_to_btn = Button(labelsFrameDownload, text='Выбрать', command=DatePickReportDateTo)
        self.date_to_btn.grid(row=10, column=3)

        ttk.Button(labelsFrameDownload, text='Выгрузить', command=self.create_docx).grid(row=11, column=2)

        # Отображение данных в окне
        # self.tree = ttk.Treeview(height=10, columns=2)
        # self.tree.grid(row=5, column=0, columnspan=2)
        # self.tree.heading('#0', text='Дата', anchor=W)
        # self.tree.heading(2, text='Статус', anchor=W)
        #
        # ttk.Button(text='Delete record').grid(row=5, column=0)
        # ttk.Button(text='Edit record').grid(row=5, column=1)
        # self.viewing_records()

    def create_docx(self):

        date_from_docx = my_gui.date_from.get()
        date_to_docx = my_gui.date_to.get()

        if self.validation_docx():

            query = "select * from reportobject where event_date >= ? and event_date <= ? order by added_date desc"
            parameters = (date_from_docx, date_to_docx,)
            query_cursor_raw = self.run_query(query, parameters)
            query_cursor = query_cursor_raw.fetchall()

            if len(query_cursor) >= 1:

                status_quantity = {}

                document = Document()

                document.add_heading('Отчет за период с ' + date_from_docx + ' по ' + date_to_docx, 2)

                for event in query_cursor:
                    print(event)
                    document.add_paragraph(
                        'Уровень: ' + event[3], style='List Bullet'
                    )
                    document.add_paragraph(
                        event[4], style='List Bullet'
                    )
                    document.add_paragraph(
                        'Срабатываний: ' + str(event[5]), style='List Bullet'
                    )
                    document.add_paragraph(
                        'Статус: ' + event[6], style='List Bullet'
                    )
                    document.add_paragraph(
                        'Комментарий: ' + event[7], style='List Bullet'
                    )

                    if event[6] in status_quantity.keys():
                        status_quantity[event[6]] += event[5]
                    else:
                        status_quantity[event[6]] = event[5]

                if 'Ложный' in status_quantity.keys():
                    document.add_paragraph(
                        'Ложных срабатываний: ' + str(status_quantity['Ложный'])
                    )
                else:
                    document.add_paragraph(
                        'Ложных срабатываний: 0'
                    )

                if 'Реальный' in status_quantity.keys():
                    document.add_paragraph(
                        'Реальных срабатываний: ' + str(status_quantity['Реальный'])
                    )
                else:
                    document.add_paragraph(
                        'Реальных срабатываний: 0'
                    )

                reports_directory = 'ReportsDocuments'

                if not os.path.exists(reports_directory):
                    os.makedirs(reports_directory)

                raw_file_name = str(datetime.datetime.now())

                new_file_name = ''

                for letter in raw_file_name:
                    if letter == ':':
                        letter = '-'
                        new_file_name += letter
                    elif letter == ' ':
                        letter = '-'
                        new_file_name += letter
                    else:
                        new_file_name += letter

                document.save(reports_directory + '/' + 'report_' + new_file_name[:-7] + '.docx')

            else:
                validation_text = 'В базе нет данных за период'
                validation_window = Toplevel()
                validation_window.title('Ошибка')
                label1 = Label(validation_window, text=validation_text, height=0, width=0, justify=LEFT, fg='red')
                label1.pack()
        else:
            validation_text = 'Укажите обе даты'
            validation_window = Toplevel()
            validation_window.title('Ошибка')
            label1 = Label(validation_window, text=validation_text, height=0, width=0, justify=LEFT, fg='red')
            label1.pack()

    def run_query(self, query, parameters=()):
        with sqlite3.connect(self.db_name) as conn:
            cursor = conn.cursor()
            query_result = cursor.execute(query, parameters)
            conn.commit()
        return query_result

    def initial_migration(self):
        query = '''create table if not exists reportobject (
                        id integer primary key,
                        event_date date,
                        added_date datetime,
                        level text,
                        event text,
                        quantity int,
                        status text,
                        comment text
                        )'''
        self.run_query(query)

    # def viewing_records(self):
    #     records = self.tree.get_children()
    #     for element in records:
    #         self.tree.delete(element)
    #     query = 'SELECT * FROM reportobject ORDER BY name DESC'
    #     db_rows = self.run_query(query)
    #     for row in db_rows:
    #         self.tree.insert('', 0, text=row[1], values=row[2])

    def validation(self):
        if (
                len(self.event_date.get()) != 0
                and len(self.level.get()) != 0
                and len(self.quantity.get()) != 0
                and self.quantity.get().isdigit()
                and len(self.status.get()) != 0
        ):
            return True

    def validation_docx(self):
        if (
                len(self.date_from.get()) != 0
                and len(self.date_to.get()) != 0
        ):
            return True

    def adding(self):
        if self.validation():
            query = '''insert into reportobject(event_date,
                                                added_date,
                                                level,
                                                event,
                                                quantity,
                                                status,
                                                comment
                                                ) values (?, ?, ?, ?, ?, ?, ?)'''
            parameters = (
                        self.event_date.get(),
                        datetime.datetime.now(),
                        self.level.get(),
                        self.event.get(),
                        self.quantity.get(),
                        self.status.get(),
                        self.comment.get('1.0', END)
                        )
            self.run_query(query, parameters)
            self.message['text'] = 'Запись добавлена'

            self.event_date.delete(0, END)
            self.event_date.insert(0, datetime.date.today())
            self.level.delete(0, END)
            self.event.delete(0, END)
            self.quantity.delete(0, END)
            self.status.delete(0, END)
            self.comment.delete('1.0', END)
        else:
            validation_text = 'Ошибка валидации данных. \n ' \
                              'Проверьте правильность заполнения формы: \n ' \
                              '1. Обязательны все поля, кроме "Комментарий" \n ' \
                              '2. Количество - только цифры.'
            validation_window = Toplevel()
            validation_window.title('Ошибка')
            label1 = Label(validation_window, text=validation_text, height=0, width=0, justify=LEFT, fg='red')
            label1.pack()

class MyDatePicker(Toplevel):
    """
    Description:
        A tkinter GUI date picker.
    """

    def __init__(self, parent=None):
        """
        Description:
            When instantiating in parent module/widget/Gui, pass in 'self' as argument.
            Ex:
                a = MyDatePicker(self)

        :param parent: parent instance.
        """

        super().__init__()
        self.parent = parent
        self.title("Календарь")
        self.resizable(0, 0)
        self.geometry("+250+10")
        self.init_frames()
        self.init_needed_vars()
        self.init_month_year_labels()
        self.init_buttons()
        self.space_between_widgets()
        self.fill_days()
        self.make_calendar()

    def init_frames(self):
        self.frame1 = Frame(self)
        self.frame1.pack()

        self.frame_days = Frame(self)
        self.frame_days.pack()

    def init_needed_vars(self):
        self.month_names = tuple(calendar.month_name)
        self.day_names = tuple(calendar.day_abbr)
        self.year = time.strftime("%Y")
        self.month = time.strftime("%B")

    def init_month_year_labels(self):
        self.year_str_var = StringVar()
        self.month_str_var = StringVar()

        self.year_str_var.set(self.year)
        self.year_lbl = Label(self.frame1, textvariable=self.year_str_var, width=3)
        self.year_lbl.grid(row=0, column=5)

        self.month_str_var.set(self.month)
        self.month_lbl = Label(self.frame1, textvariable=self.month_str_var, width=8)
        self.month_lbl.grid(row=0, column=1)

    def init_buttons(self):
        self.left_yr = ttk.Button(self.frame1, text="←", width=5, command=self.prev_year)
        self.left_yr.grid(row=0, column=4)

        self.right_yr = ttk.Button(self.frame1, text="→", width=5, command=self.next_year)
        self.right_yr.grid(row=0, column=6)

        self.left_mon = ttk.Button(self.frame1, text="←", width=5, command=self.prev_month)
        self.left_mon.grid(row=0, column=0)

        self.right_mon = ttk.Button(self.frame1, text="→", width=5, command=self.next_month)
        self.right_mon.grid(row=0, column=2)

    def space_between_widgets(self):
        self.frame1.grid_columnconfigure(3, minsize=40)

    def prev_year(self):
        self.prev_yr = int(self.year_str_var.get()) - 1
        self.year_str_var.set(self.prev_yr)

        self.make_calendar()

    def next_year(self):
        self.next_yr = int(self.year_str_var.get()) + 1
        self.year_str_var.set(self.next_yr)

        self.make_calendar()

    def prev_month(self):
        index_current_month = int(self.month_names.index(self.month_str_var.get()))
        index_prev_month = index_current_month - 1

        #  index 0 is empty string, use index 12 instead, which is index of December.
        if index_prev_month == 0:
            self.month_str_var.set(self.month_names[12])
        else:
            self.month_str_var.set(self.month_names[index_current_month - 1])

        self.make_calendar()

    def next_month(self):
        index_current_month = int(self.month_names.index(self.month_str_var.get()))

        #  index 13 does not exist, use index 1 instead, which is January.
        try:
            self.month_str_var.set(self.month_names[index_current_month + 1])
        except IndexError:
            self.month_str_var.set(self.month_names[1])

        self.make_calendar()

    def fill_days(self):
        col = 0
        #  Creates days label
        for day in self.day_names:
            self.lbl_day = Label(self.frame_days, text=day)
            self.lbl_day.grid(row=0, column=col)
            col += 1

    def make_calendar(self):
        #  Delete date buttons if already present.
        #  Each button must have its own instance attribute for this to work.
        try:
            for dates in self.m_cal:
                for date in dates:
                    if date == 0:
                        continue

                    self.delete_buttons(date)

        except AttributeError:
            pass

        year = int(self.year_str_var.get())
        month = self.month_names.index(self.month_str_var.get())
        self.m_cal = calendar.monthcalendar(year, month)

        #  build date buttons.
        for dates in self.m_cal:
            row = self.m_cal.index(dates) + 1
            for date in dates:
                col = dates.index(date)

                if date == 0:
                    continue

                self.make_button(str(date), str(row), str(col))

    def make_button(self, date, row, column):
        exec(
            "self.btn_" + date + "= ttk.Button(self.frame_days, text=" + date + ", width=5)\n"
                                                                                "self.btn_" + date + ".grid(row=" + row + " , column=" + column + ")\n"
                                                                                                                                                  "self.btn_" + date + ".bind(\"<Button-1>\", self.get_date)"
        )

    def delete_buttons(self, date):
        exec(
            "self.btn_" + str(date) + ".destroy()"
        )

    def get_date(self, clicked=None):
        clicked_button = clicked.widget
        year = self.year_str_var.get()
        month = self.month_names.index(self.month_str_var.get())
        date = clicked_button['text']

        #  Change string format for different date formats.
        self.full_date = '%s-%02d-%02d' % (year, month, date)


class DatePickEventDate(MyDatePicker):

    def get_date(self, clicked=None):
        clicked_button = clicked.widget
        year = self.year_str_var.get()
        month = self.month_names.index(self.month_str_var.get())
        date = clicked_button['text']
        self.full_date = '%s-%02d-%02d' % (year, month, date)
        my_gui.event_date.delete(0, END)
        my_gui.event_date.insert(0, self.full_date)
        DatePickEventDate.destroy(self)


class DatePickReportDateFrom(MyDatePicker):

    def get_date(self, clicked=None):
        clicked_button = clicked.widget
        year = self.year_str_var.get()
        month = self.month_names.index(self.month_str_var.get())
        date = clicked_button['text']
        self.full_date = '%s-%02d-%02d' % (year, month, date)
        my_gui.date_from.delete(0, END)
        my_gui.date_from.insert(0, self.full_date)
        DatePickReportDateFrom.destroy(self)


class DatePickReportDateTo(MyDatePicker):

    def get_date(self, clicked=None):
        clicked_button = clicked.widget
        year = self.year_str_var.get()
        month = self.month_names.index(self.month_str_var.get())
        date = clicked_button['text']
        self.full_date = '%s-%02d-%02d' % (year, month, date)
        my_gui.date_to.delete(0, END)
        my_gui.date_to.insert(0, self.full_date)
        DatePickReportDateTo.destroy(self)


if __name__ == '__main__':

    root = Tk()
    my_gui = ReportObject(root)
    root.mainloop()